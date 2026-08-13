from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from whatsapp_travel_catalogue_data import SECTIONS as TRAVEL_SECTIONS


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "easyTag-whatsapp-message-approval-catalogue.docx"

INK = "17212B"
GREEN = "169B62"
GREEN_DARK = "0E7147"
GREEN_PALE = "EAF7F1"
CORAL = "E56854"
CORAL_PALE = "FFF1EE"
BLUE_PALE = "EDF5FB"
GRAY = "66717D"
GRAY_PALE = "F4F6F7"
LINE = "DCE3E6"
WHITE = "FFFFFF"


LEGACY_SECTIONS = [
    {
        "title": "New contacts who have not purchased",
        "intro": "For eligible contacts with a phone number but no order history. Use broad human moments; do not pretend to know what they lose or what device they own.",
        "items": [
            {
                "title": "The wallet self-test",
                "when": "No purchase; no confirmed product interest; broad Apple or Samsung-compatible destination available.",
                "message": "Tiny question: where is your wallet right now?\n\nIf that took even a second, you will understand why this exists.",
                "button": "See why",
                "destination": "A mobile-first wallet tracker page that identifies Apple and Samsung choices immediately.",
                "lever": "A tiny real-world action creates instant personal relevance.",
                "guardrail": "Do not send to someone who already owns the relevant wallet tracker.",
            },
            {
                "title": "The same-pocket loop",
                "when": "No purchase; suitable when wallet products are available for the recipient's confirmed ecosystem.",
                "message": "Be honest: how many times do you check the same pocket when your wallet goes missing?\n\nThere is a card made for shortening that particular little drama.",
                "button": "See the card",
                "destination": "The matching wallet card product page, with the card visible above the fold.",
                "lever": "Recognition and gentle humour, without creating fear.",
                "guardrail": "The linked card must be compatible with the recipient's phone ecosystem.",
            },
            {
                "title": "The room that swallowed it",
                "when": "No purchase; broad acquisition message for a general tracker guide or collection page.",
                "message": "Nothing is truly lost until the room has watched you search the same three places twice.\n\nThere is a calmer way to do this.",
                "button": "Show me",
                "destination": "A short guide showing how easyTag helps locate everyday items.",
                "lever": "Familiar frustration followed by relief.",
                "guardrail": "Avoid unsupported range or accuracy claims on the destination.",
            },
            {
                "title": "The miniature treasure hunt",
                "when": "No purchase; useful for prospects who have shown interest in keys or everyday carry categories.",
                "message": "Keys have a strange talent for turning an ordinary morning into a miniature treasure hunt.\n\nThis makes the ending much less dramatic.",
                "button": "End the hunt",
                "destination": "A key-tracking product or use-case page, if that product is currently sold.",
                "lever": "Humour plus a clear promise of reduced hassle.",
                "guardrail": "Only activate after the current key-tracker product and landing page are verified.",
            },
            {
                "title": "The compatibility shortcut",
                "when": "No purchase; phone ecosystem is unknown or the person previously browsed more than one tracker type.",
                "message": "Apple or Samsung? That one detail changes which tracker makes sense.\n\nWe made the answer pleasantly quick.",
                "button": "Find mine",
                "destination": "A two-choice compatibility selector, not a generic catalogue page.",
                "lever": "Removes uncertainty with a low-effort next step.",
                "guardrail": "The selector must answer the question in one tap and show only valid products.",
            },
            {
                "title": "The leaving-the-house check",
                "when": "No purchase; general everyday-carry audience.",
                "message": "Phone. Keys. Wallet.\n\nThe three-second check before leaving home has a surprisingly useful upgrade.",
                "button": "See the upgrade",
                "destination": "An editorial landing page connecting the three-item ritual to one relevant tracker.",
                "lever": "A universal routine creates rhythm and curiosity.",
                "guardrail": "Do not turn the destination into a multi-product wall; lead with one use case.",
            },
            {
                "title": "The future-you favour",
                "when": "No purchase; broad prospecting when a clear everyday item use case is available.",
                "message": "The best time to make something findable is before it disappears.\n\nFuture you tends to appreciate that sort of planning.",
                "button": "Do future me a favour",
                "destination": "A concise use-case page with a clear Apple/Samsung route.",
                "lever": "Positive anticipation rather than loss anxiety.",
                "guardrail": "Keep the page practical; avoid alarmist lost-item stories.",
            },
            {
                "title": "The surprising first choice",
                "when": "No purchase; suitable for an educational quiz or guide rather than a direct product page.",
                "message": "People rarely agree on the first thing worth tracking.\n\nThe answers are more revealing than you might expect.",
                "button": "See the list",
                "destination": "A short editorial list of real use cases, ending in ecosystem-matched products.",
                "lever": "Curiosity and self-comparison.",
                "guardrail": "The page must contain a genuinely useful list, not disguised catalogue copy.",
            },
        ],
    },
    {
        "title": "Previous AI Inbox contacts after resolution",
        "intro": "For people who contacted the business and whose support issue is fully resolved. Their conversation makes them eligible, but the marketing message should not expose or quote private support details.",
        "items": [
            {
                "title": "The practical question",
                "when": "A prior inbox conversation is resolved; at least a seven-day cooling period; no open complaint or refund.",
                "message": "One practical question we hear a lot: what is actually worth making findable?\n\nWe put the useful answers in one place.",
                "button": "See the answers",
                "destination": "A helpful use-case guide matched to known ecosystem where possible.",
                "lever": "Feels like useful follow-up without referring to the person's private issue.",
                "guardrail": "Never send while their conversation is unresolved, escalated, refunded, or negative.",
            },
            {
                "title": "The one-minute guide",
                "when": "Previous contact; issue resolved; recipient has not purchased.",
                "message": "We made a one-minute guide for choosing a tracker without reading twelve nearly identical product pages.\n\nIt is refreshingly short.",
                "button": "Read the short version",
                "destination": "A one-minute compatibility and use-case guide.",
                "lever": "Promises relief from decision fatigue.",
                "guardrail": "The guide must genuinely be short and current.",
            },
            {
                "title": "The overlooked use",
                "when": "Previous contact; issue resolved; no recent marketing message using the same product category.",
                "message": "There is one easyTag use people tend to overlook until they see it.\n\nIt is not the obvious one.",
                "button": "What is it?",
                "destination": "An editorial page that reveals one specific, credible use above the fold.",
                "lever": "A clean open loop with a promised answer.",
                "guardrail": "The reveal must be meaningful; reject vague or underwhelming destinations.",
            },
            {
                "title": "The phone-first answer",
                "when": "Previous contact; ecosystem is known but no purchase has been made.",
                "message": "The right tracker is mostly a phone question, not a product question.\n\nOnce that is clear, the choice gets much easier.",
                "button": "See my match",
                "destination": "A pre-filtered Apple or Samsung recommendation page.",
                "lever": "Reframes a confusing purchase into one simple decision.",
                "guardrail": "Only pre-filter when ecosystem data is reliable and recently confirmed.",
            },
            {
                "title": "The useful rabbit hole",
                "when": "Previous contact; resolved; audience has engaged with guides or help content.",
                "message": "We found a surprisingly useful rabbit hole: the ordinary things people most regret not tracking.\n\nThis is the short version.",
                "button": "Take a look",
                "destination": "A concise editorial list or quiz with a relevant product bridge.",
                "lever": "Curiosity with a conversational, non-corporate tone.",
                "guardrail": "Avoid stories that exploit loss, theft, or personal distress.",
            },
            {
                "title": "The no-hard-sell note",
                "when": "Previous contact; resolved; useful as a softer re-entry after a longer cooling period.",
                "message": "No grand announcement. We just made it easier to work out which easyTag fits which phone and item.\n\nThat seemed worth sharing.",
                "button": "Use the matcher",
                "destination": "The compatibility matcher.",
                "lever": "Disarms advertising expectations through understatement.",
                "guardrail": "Do not use if the matcher is incomplete or requires a long form.",
            },
        ],
    },
    {
        "title": "First-time buyers and new owners",
        "intro": "For customers after delivery or confirmed activation. These should increase product value first; only introduce another product when the first experience is healthy and the second use case is genuinely complementary.",
        "items": [
            {
                "title": "The first test run",
                "when": "Order delivered or activation confirmed; ideally 3-7 days later.",
                "message": "A tiny suggestion: test your easyTag once while you already know where it is.\n\nIt makes the real moment much less mysterious.",
                "button": "Do the test",
                "destination": "A short ecosystem-specific test guide.",
                "lever": "Helpful ownership advice with immediate utility.",
                "guardrail": "Use the correct Apple or Samsung instructions; no cross-sell on this page.",
            },
            {
                "title": "The setting people skip",
                "when": "New owner; activation is known or likely complete; no unresolved setup ticket.",
                "message": "There is one setup check people often skip because everything already seems to work.\n\nIt takes about thirty seconds.",
                "button": "Check mine",
                "destination": "The exact setup check, visible immediately.",
                "lever": "Useful curiosity without implying their product is faulty.",
                "guardrail": "Do not invent a setting; product support must verify the step and timing claim.",
            },
            {
                "title": "The best first job",
                "when": "New owner; product not yet associated with a known item.",
                "message": "The tracker is ready. The harder question is what deserves it first.\n\nWe made a very short shortlist.",
                "button": "Choose its job",
                "destination": "A use-case chooser appropriate to the purchased model.",
                "lever": "Turns setup into a small personal decision.",
                "guardrail": "Exclude use cases the purchased model cannot support.",
            },
            {
                "title": "The boring-but-brilliant habit",
                "when": "New owner; 7-14 days after delivery; support state healthy.",
                "message": "The most useful tracker habit is also the least exciting: check it before the day you need it.\n\nHere is the ten-second version.",
                "button": "Show the habit",
                "destination": "A lightweight owner-care or check-in guide.",
                "lever": "Honesty and practical value build trust.",
                "guardrail": "Any maintenance guidance must match the exact product model.",
            },
            {
                "title": "The second-item test",
                "when": "First product is delivered and no issue is open; customer owns only one tracked-item category.",
                "message": "Quick test: what is the next thing you would turn around to retrieve?\n\nThat answer is usually more useful than a product recommendation.",
                "button": "Try the test",
                "destination": "A one-question cross-category selector.",
                "lever": "Self-discovery makes a cross-sell feel earned.",
                "guardrail": "Do not show the product they already own as the primary recommendation.",
            },
            {
                "title": "The near-miss memory",
                "when": "New owner; 14+ days after delivery; suitable for cross-category discovery.",
                "message": "Think of the last thing you nearly left behind.\n\nThere may be a very small fix for that exact feeling.",
                "button": "See the options",
                "destination": "A use-case page filtered to products the customer does not own.",
                "lever": "A personal memory creates relevance without pressure.",
                "guardrail": "Keep language light; do not imply access to their location or behaviour.",
            },
            {
                "title": "The owner shortcut",
                "when": "New owner; known product; no recent support issue.",
                "message": "You should not need to become a tracker expert to get the useful bits right.\n\nWe condensed them into one page.",
                "button": "Open the shortcut",
                "destination": "A model-specific owner cheat sheet.",
                "lever": "Competence and convenience, not selling.",
                "guardrail": "The cheat sheet must be kept in sync with current app instructions.",
            },
            {
                "title": "The oddly satisfying check",
                "when": "New owner; engagement message after successful delivery.",
                "message": "There is something oddly satisfying about making the tracker ring while it is sitting right beside you.\n\nConsider this your official test excuse.",
                "button": "Test it now",
                "destination": "The correct ring/test instructions for the purchased product.",
                "lever": "Playfulness encourages useful product engagement.",
                "guardrail": "Only use if the purchased model supports the described ring action.",
            },
        ],
    },
    {
        "title": "Confirmed Apple users",
        "intro": "For customers whose Apple ecosystem is known from a purchase, setup record, or explicit conversation. These routes may reference Find My, but every technical claim still needs product-page verification.",
        "items": [
            {
                "title": "The app already waiting",
                "when": "Confirmed Apple user; does not already own the featured Apple-compatible product.",
                "message": "The useful part may already be on your iPhone.\n\nThe interesting part is what you can make visible inside Find My.",
                "button": "See what fits",
                "destination": "An Apple-compatible collection or selector led by Find My use cases.",
                "lever": "Uses familiar technology to lower perceived complexity.",
                "guardrail": "Verify exact Find My compatibility for every product shown.",
            },
            {
                "title": "The wallet-shaped answer",
                "when": "Confirmed Apple user; no wallet card purchase.",
                "message": "A tracker is useful. A tracker shaped like the thing it protects is rather better.\n\nThis one disappears into a wallet.",
                "button": "See the fit",
                "destination": "Apple-compatible wallet card page with thickness and fit shown honestly.",
                "lever": "A satisfying form-factor reveal.",
                "guardrail": "Do not make unverified thickness or universal-wallet claims in copy.",
            },
            {
                "title": "The Find My thought experiment",
                "when": "Confirmed Apple user; broad product discovery.",
                "message": "Open Find My for a second.\n\nWhat is the one thing you wish appeared there but does not yet?",
                "button": "Find the right shape",
                "destination": "An Apple product selector organized by item shape and use case.",
                "lever": "Turns the phone into part of the message interaction.",
                "guardrail": "Do not ask this if the destination cannot answer by item type quickly.",
            },
            {
                "title": "The quiet travel upgrade",
                "when": "Confirmed Apple user; travel interest or relevant product browsing; no matching purchase.",
                "message": "The nicest travel upgrades are often the ones you barely notice.\n\nUntil the moment you are very glad they are there.",
                "button": "See the small one",
                "destination": "A single Apple-compatible travel product or focused travel page.",
                "lever": "Understatement and anticipation.",
                "guardrail": "Avoid implying guaranteed recovery, airline tracking, or theft prevention.",
            },
            {
                "title": "The card-versus-tag choice",
                "when": "Confirmed Apple user comparing product formats.",
                "message": "Card or tag? It is less about the technology and more about where it has to live.\n\nThe thirty-second answer is here.",
                "button": "Choose the shape",
                "destination": "A side-by-side format guide for Apple-compatible products.",
                "lever": "Simplifies a real decision without hard selling.",
                "guardrail": "Comparison facts and available models must be current.",
            },
            {
                "title": "The invisible everyday upgrade",
                "when": "Confirmed Apple user; suitable for a slim or discreet product format.",
                "message": "A good everyday upgrade should mostly disappear.\n\nThis one sits quietly until you need Find My to notice it.",
                "button": "See where it fits",
                "destination": "A focused Apple-compatible product page showing the real object in use.",
                "lever": "Signals simplicity and low friction.",
                "guardrail": "Use only for a product that genuinely fits the depicted use case.",
            },
        ],
    },
    {
        "title": "Confirmed Samsung users",
        "intro": "For customers whose Samsung ecosystem is reliably known. This group deserves Samsung-native language rather than an Apple-first comparison, except where the contrast itself is the useful surprise.",
        "items": [
            {
                "title": "The Samsung network surprise",
                "when": "Confirmed Samsung user; no Samsung wallet tracker purchase.",
                "message": "AirTags get all the attention, but Samsung phones have their own finding network.\n\nThe wallet-sized version is worth seeing.",
                "button": "Show me",
                "destination": "The Samsung-compatible wallet tracker page.",
                "lever": "Corrects a common awareness gap with a useful surprise.",
                "guardrail": "Verify current SmartThings Find naming and compatibility before activation.",
            },
            {
                "title": "The no-Apple-workaround note",
                "when": "Confirmed Samsung user; suitable for broad Samsung product discovery.",
                "message": "Samsung users should not need an Apple workaround to make everyday things findable.\n\nThere is a version built for your side of the phone aisle.",
                "button": "See the Samsung one",
                "destination": "A Samsung-specific landing page, not a mixed catalogue.",
                "lever": "Recognition and ecosystem belonging.",
                "guardrail": "Keep the tone friendly; do not disparage Apple or overstate exclusivity.",
            },
            {
                "title": "The SmartThings question",
                "when": "Confirmed Samsung user; has not purchased the featured tracker.",
                "message": "What would you add to SmartThings Find if it could be almost anything you carry?\n\nThe wallet answer is unexpectedly neat.",
                "button": "See the wallet answer",
                "destination": "The Samsung wallet product page with setup path visible.",
                "lever": "Invites imagination, then provides one concrete reveal.",
                "guardrail": "Only use SmartThings language confirmed by current product instructions.",
            },
            {
                "title": "The Samsung compatibility shortcut",
                "when": "Confirmed Samsung user comparing available tracker formats.",
                "message": "The words 'works with Android' are not specific enough.\n\nHere is the short list made for Samsung phones.",
                "button": "Open the short list",
                "destination": "A verified Samsung compatibility guide.",
                "lever": "Calls out a real source of confusion and resolves it.",
                "guardrail": "Every device and app requirement on the page must be explicit and current.",
            },
            {
                "title": "The wallet demonstration",
                "when": "Confirmed Samsung user; wallet product available; no existing wallet tracker purchase.",
                "message": "The clever bit is not that it can be found.\n\nIt is that a Samsung-compatible tracker can look like an ordinary card.",
                "button": "Watch it disappear",
                "destination": "A product page or short demo showing the card placed in a real wallet.",
                "lever": "Visual curiosity around the product's form.",
                "guardrail": "The destination must show the promised demonstration above the fold.",
            },
            {
                "title": "The Samsung owner's quick route",
                "when": "Confirmed Samsung user; broad acquisition or cross-category discovery.",
                "message": "We removed the Apple options, the vague compatibility claims and the guesswork.\n\nThis is the Samsung-only route.",
                "button": "Take the quick route",
                "destination": "A Samsung-filtered product and use-case page.",
                "lever": "Clarity and respect for the recipient's ecosystem.",
                "guardrail": "Do not send if the Samsung catalogue is incomplete or the filter can reset.",
            },
        ],
    },
    {
        "title": "Travel, luggage and passport scenarios",
        "intro": "For customers with a travel-category purchase, browsing signal, or relevant prior enquiry. Use familiar travel moments, never exaggerated danger or promises that a tracker replaces airline, customs, or security procedures.",
        "items": [
            {
                "title": "The conveyor-belt trust exercise",
                "when": "Travel interest; no matching luggage tracker purchase.",
                "message": "A suitcase disappearing behind the conveyor belt is a surprisingly large trust exercise.\n\nThere is a small way to make it feel less blind.",
                "button": "See how",
                "destination": "A luggage tracking use-case page for the confirmed phone ecosystem.",
                "lever": "A vivid shared moment and a proportionate promise.",
                "guardrail": "Do not imply live airline data, guaranteed location, or recovery.",
            },
            {
                "title": "The airport-pocket divide",
                "when": "Travel interest; passport or wallet format available.",
                "message": "Airport security has two groups: the organised, and the suddenly-checking-every-pocket.\n\nWe made something for the second group.",
                "button": "That is me",
                "destination": "A passport or wallet tracking page that matches the hook exactly.",
                "lever": "Identity, recognition and gentle humour.",
                "guardrail": "Choose one product format; do not send people to a mixed travel collection.",
            },
            {
                "title": "The hotel-room sweep",
                "when": "Travel interest; broad travel accessory discovery.",
                "message": "The final hotel-room sweep always feels suspiciously incomplete.\n\nThere is one way to make the last look less theatrical.",
                "button": "Do the calmer check",
                "destination": "A practical travel checklist featuring one relevant easyTag use.",
                "lever": "A highly recognizable scene with light humour.",
                "guardrail": "The checklist must provide real value beyond the product link.",
            },
            {
                "title": "The passport-sleeve reveal",
                "when": "Travel interest; compatible passport tracker cover is currently available; no existing purchase.",
                "message": "A passport cover can do more than protect the corners.\n\nThe useful part is hidden inside this one.",
                "button": "Open the cover",
                "destination": "The passport tracker cover page with the tracking feature demonstrated immediately.",
                "lever": "A physical reveal and clear curiosity gap.",
                "guardrail": "Verify product availability, passport fit, and ecosystem compatibility.",
            },
            {
                "title": "The gate-number ritual",
                "when": "Travel interest; educational or checklist destination available.",
                "message": "You can read the gate number three times and still check it again five minutes later.\n\nTravel does strange things to otherwise sensible people.",
                "button": "See the useful list",
                "destination": "A concise pre-travel checklist with a relevant tracker use case.",
                "lever": "Warm recognition builds affinity before the product appears.",
                "guardrail": "The landing page must continue the same conversational tone.",
            },
            {
                "title": "The bag that looks exactly like yours",
                "when": "Travel interest; luggage category available.",
                "message": "Every baggage belt contains at least six bags that look exactly like yours from twenty metres away.\n\nOne small addition makes the waiting less vague.",
                "button": "See the addition",
                "destination": "A luggage tracking or identification page matched to ecosystem.",
                "lever": "Visual recognition and mild uncertainty.",
                "guardrail": "Avoid claiming the tracker visually identifies the bag unless the product truly does.",
            },
        ],
    },
    {
        "title": "Repeat buyers and lapsed customers",
        "intro": "For customers with a healthy prior purchase history. Use what is known to avoid duplicate recommendations, but do not expose order detail in the message itself unless a later personalization test explicitly approves it.",
        "items": [
            {
                "title": "The second problem",
                "when": "Existing owner; first product delivered; no unresolved issue; another relevant category remains unowned.",
                "message": "You have already solved one 'where did I put it?' problem.\n\nThe second one is usually easier to identify than expected.",
                "button": "Find the second one",
                "destination": "A cross-category selector excluding products already purchased.",
                "lever": "Builds on success without naming private order details.",
                "guardrail": "Purchase exclusions must be reliable and account for variants or duplicate SKUs.",
            },
            {
                "title": "The different reason",
                "when": "Repeat buyer or engaged owner; 30+ days after first healthy purchase.",
                "message": "Most people choose their second tracker for a completely different reason than the first.\n\nThe pattern is surprisingly consistent.",
                "button": "See the pattern",
                "destination": "An editorial page showing credible second-use scenarios.",
                "lever": "Social curiosity without a popularity claim that requires exact numbers.",
                "guardrail": "Phrase as an editorial observation unless real purchase data substantiates it.",
            },
            {
                "title": "The next turn-around item",
                "when": "Existing owner; no recent marketing; complementary categories available.",
                "message": "What is the next thing that would make you turn around halfway down the street?\n\nThat is probably the useful answer.",
                "button": "Match the answer",
                "destination": "A quick use-case matcher excluding owned products.",
                "lever": "A vivid self-test produces personal relevance.",
                "guardrail": "Do not infer or state what the answer is from private data.",
            },
            {
                "title": "The quiet return",
                "when": "Lapsed customer; no purchase or inbox activity for at least 90 days; previous experience healthy.",
                "message": "It has been a while, so we will skip the grand comeback speech.\n\nThere is simply a useful new route through the easyTag range.",
                "button": "Take the quick look",
                "destination": "A streamlined product matcher or genuinely updated collection.",
                "lever": "Acknowledges the gap without sounding needy or promotional.",
                "guardrail": "Only say 'new' if the route, range, or experience is genuinely new to that customer.",
            },
            {
                "title": "The owner-only shortcut",
                "when": "Existing owner; a useful owner resource or complementary selector is available.",
                "message": "This is more useful if you already own an easyTag.\n\nIt starts where the normal product page finishes.",
                "button": "Open the owner bit",
                "destination": "An owner-focused guide, care page, or advanced use-case selector.",
                "lever": "Insider relevance without using fake exclusivity.",
                "guardrail": "The destination must genuinely differ from a normal product page.",
            },
            {
                "title": "The one-that-earned-its-place",
                "when": "Repeat buyer; product engagement or positive ownership is known; no active support issue.",
                "message": "The best test of a small gadget is whether it quietly earns its place.\n\nIf yours has, this next use might make sense too.",
                "button": "See the next use",
                "destination": "One complementary use case, chosen from purchase exclusions.",
                "lever": "Respectful acknowledgement of product value.",
                "guardrail": "Do not send where satisfaction is unknown after a complaint, return, or failed setup.",
            },
        ],
    },
    {
        "title": "Seasonal and situational moments",
        "intro": "For broad cohorts when the timing is genuinely relevant. These concepts should be selected by real season, travel period, gift window, or routine change, not used as artificial urgency.",
        "items": [
            {
                "title": "The pre-trip ten minutes",
                "when": "Travel season or known travel interest; send early enough to be useful, not as last-minute pressure.",
                "message": "Ten quiet minutes before a trip can save a lot of frantic pocket-checking later.\n\nThis is the useful version of that checklist.",
                "button": "Open the checklist",
                "destination": "A concise pre-trip checklist with ecosystem-matched product routes.",
                "lever": "Preparedness and relief.",
                "guardrail": "No countdowns, false deadlines, or unverified delivery promises.",
            },
            {
                "title": "The gift with a future story",
                "when": "Gifting season; recipient is eligible; product giftability and delivery timing are confirmed.",
                "message": "Some gifts are appreciated immediately. Others become brilliant the first time they prevent a small disaster.\n\nThis is the second kind.",
                "button": "See the idea",
                "destination": "One giftable easyTag use case with clear compatibility guidance.",
                "lever": "Future storytelling and emotional payoff.",
                "guardrail": "Avoid guaranteed rescue language and confirm dispatch information before use.",
            },
            {
                "title": "The routine reset",
                "when": "Back-to-work, school, or routine-change period; relevant everyday carry destination exists.",
                "message": "New routines are excellent at revealing which small things never had a proper home.\n\nThere is a neat fix for the most mobile ones.",
                "button": "See the neat fix",
                "destination": "An everyday-carry use-case page matched to ecosystem.",
                "lever": "Timely recognition without a sale event.",
                "guardrail": "Do not infer children, school status, or employment from weak data.",
            },
            {
                "title": "The Sunday drawer test",
                "when": "General engagement; suitable for a guide or quiz; avoid high-frequency use.",
                "message": "A Sunday drawer tidy has one awkward question: why do we own three cables and still lose the keys?\n\nOne of those problems has a simpler answer.",
                "button": "See the answer",
                "destination": "A light editorial page focused on the relevant trackable item.",
                "lever": "Domestic humour and a satisfying contrast.",
                "guardrail": "Keep it occasional; the joke will wear out quickly if repeated.",
            },
        ],
    },
]

SECTIONS = TRAVEL_SECTIONS


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=LINE, size=6):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep(paragraph, keep_next=False):
    p_pr = paragraph._p.get_or_add_pPr()
    keep_lines = OxmlElement("w:keepLines")
    p_pr.append(keep_lines)
    if keep_next:
        node = OxmlElement("w:keepNext")
        p_pr.append(node)


def set_paragraph_box(paragraph, fill, border_color):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "5")
        node.set(qn("w:color"), border_color)
        borders.append(node)
    p_pr.append(borders)


def set_font(run, size=10.5, color=INK, bold=False, italic=False, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_text(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    set_font(run, **kwargs)
    return run


def set_spacing(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(paragraph, "PAGE ", size=8.5, color=GRAY, bold=True)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_font(run, size=8.5, color=GRAY, bold=True)


def add_label_value(doc, label, value, after=2):
    p = doc.add_paragraph()
    set_spacing(p, after=after, line=1.08)
    add_text(p, f"{label}: ", size=9.3, color=GREEN_DARK, bold=True)
    add_text(p, value, size=9.3, color=INK)
    return p


def add_message_box(doc, message, button):
    p = doc.add_paragraph()
    set_spacing(p, before=1, after=4, line=1.12)
    set_keep(p, keep_next=True)
    set_paragraph_box(p, GREEN_PALE, "B9DDCC")
    chunks = message.split("\n\n")
    for idx, chunk in enumerate(chunks):
        if idx:
            p.add_run().add_break()
            p.add_run().add_break()
        add_text(p, chunk, size=10.3, color=INK)
    button_p = doc.add_paragraph()
    set_spacing(button_p, before=0, after=4)
    set_keep(button_p, keep_next=True)
    add_text(button_p, "BUTTON COPY  ", size=8.8, color=GREEN_DARK, bold=True)
    add_text(button_p, button, size=9.5, color=INK, bold=True)


def add_option(doc, number, item):
    heading = doc.add_paragraph()
    set_spacing(heading, before=9, after=4)
    set_keep(heading, keep_next=True)
    add_text(heading, f"OPTION {number:02d}", size=9, color=CORAL, bold=True)
    add_text(heading, f"  |  {item['title']}", size=12.2, color=INK, bold=True)

    use_p = add_label_value(doc, "Use when", item["when"], after=5)
    set_keep(use_p, keep_next=True)
    add_message_box(doc, item["message"], item["button"])

    add_label_value(doc, "Destination", item["destination"])
    add_label_value(doc, "Why it may earn the click", item["lever"])
    add_label_value(doc, "Guardrail", item["guardrail"], after=3)

    decision = doc.add_paragraph()
    set_spacing(decision, before=2, after=7)
    add_text(decision, "DECISION   [ ] APPROVE   [ ] REVISE   [ ] REJECT", size=8.8, color=GRAY, bold=True)
    p_pr = decision._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), LINE)
    border.append(bottom)
    p_pr.append(border)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.36)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 30, INK, 0, 8),
        ("Subtitle", 13, GRAY, 0, 18),
        ("Heading 1", 18, INK, 16, 8),
        ("Heading 2", 13, GREEN_DARK, 12, 6),
        ("Heading 3", 11.5, INK, 9, 4),
    ):
        style = styles[name]
        style.font.name = "Aptos Display" if name in ("Title", "Heading 1") else "Aptos"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    set_spacing(hp, after=0)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(hp, "easyTag", size=9, color=GREEN, bold=True)
    add_text(hp, "  /  WhatsApp creative approval", size=8.5, color=GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)

    kicker = doc.add_paragraph()
    set_spacing(kicker, before=34, after=5)
    add_text(kicker, "CREATIVE APPROVAL CATALOGUE", size=9.5, color=CORAL, bold=True)

    title = doc.add_paragraph(style="Title")
    set_spacing(title, after=8)
    add_text(title, "WhatsApp messages that do not feel like marketing", size=30, color=INK, bold=True, name="Aptos Display")

    subtitle = doc.add_paragraph(style="Subtitle")
    add_text(subtitle, "50 travel-led concepts built around flying, airports, holidays and the easy family trust connection.", size=13, color=GRAY)

    rule = doc.add_paragraph()
    set_spacing(rule, after=16)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "16")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), GREEN)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    p = doc.add_paragraph()
    set_spacing(p, after=0, line=1.2)
    set_paragraph_box(p, CORAL_PALE, "F3C7BF")
    add_text(p, "HOW TO USE THIS DOCUMENT", size=9.5, color=CORAL, bold=True)
    p.add_run().add_break()
    p.add_run().add_break()
    add_text(p, "Reply with the option numbers you approve, reject, or want revised. Copy approval does not approve an audience, date, spend, template submission, or send. Nothing in this document is scheduled.", size=10.2, color=INK)

    doc.add_heading("Approval principles", level=1)
    principles = [
        "One message, one idea, one click. The website handles conversion.",
        "The destination must satisfy the curiosity created in the message immediately on mobile.",
        "Known purchase, ecosystem and support data should improve relevance, never expose private detail.",
        "No open support issue, complaint, return or refund enters a marketing audience.",
        "Every final campaign still requires audience, timing, landing page and Meta template approval.",
    ]
    for text_value in principles:
        p = doc.add_paragraph(style="List Bullet")
        set_spacing(p, after=4, line=1.18)
        add_text(p, text_value, size=10.3, color=INK)

    doc.add_heading("Catalogue map", level=1)
    summary = doc.add_table(rows=1, cols=3)
    summary.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(summary, [820, 3440, 5100], 130)
    headers = ["Options", "Audience", "Primary role"]
    for idx, value in enumerate(headers):
        cell = summary.rows[0].cells[idx]
        set_cell_shading(cell, INK)
        set_cell_border(cell, color=INK, size=6)
        set_cell_margins(cell, top=100, start=130, bottom=100, end=130)
        p = cell.paragraphs[0]
        set_spacing(p, after=0)
        add_text(p, value, size=9, color=WHITE, bold=True)
    set_repeat_table_header(summary.rows[0])
    start = 1
    for section_data in SECTIONS:
        end = start + len(section_data["items"]) - 1
        row = summary.add_row()
        values = [f"{start:02d}-{end:02d}", section_data["title"], section_data["intro"].split(".")[0] + "."]
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            set_cell_shading(cell, WHITE if start % 2 else GRAY_PALE)
            set_cell_border(cell, color=LINE, size=5)
            set_cell_margins(cell, top=95, start=130, bottom=95, end=130)
            p = cell.paragraphs[0]
            set_spacing(p, after=0, line=1.08)
            add_text(p, value, size=8.8, color=INK, bold=idx == 0)
        start = end + 1

    doc.add_heading("Approval record", level=1)
    for label in ("Approve", "Revise", "Reject", "Notes"):
        p = doc.add_paragraph()
        set_spacing(p, after=8)
        add_text(p, f"{label}: ", size=10, color=GREEN_DARK, bold=True)
        add_text(p, "____________________________________________________________", size=10, color=LINE)

    doc.add_page_break()

    number = 1
    for section_index, section_data in enumerate(SECTIONS):
        h = doc.add_heading(section_data["title"], level=1)
        set_keep(h, keep_next=True)
        intro = doc.add_paragraph()
        set_spacing(intro, after=9, line=1.18)
        set_keep(intro, keep_next=True)
        add_text(intro, section_data["intro"], size=10.2, color=GRAY, italic=True)
        for item in section_data["items"]:
            add_option(doc, number, item)
            number += 1
        if section_index < len(SECTIONS) - 1:
            doc.add_page_break()

    doc.add_page_break()
    doc.add_heading("Recommended first shortlist", level=1)
    p = doc.add_paragraph()
    set_spacing(p, after=9, line=1.22)
    add_text(p, "These are the strongest starting directions for a copy discussion. This is not send approval.", size=10.5, color=INK)

    shortlist = [
        (1, "Best broad flying-soon opener for non-buyers."),
        (4, "Best emotionally recognisable check-in moment."),
        (9, "Best humorous airport identity hook."),
        (17, "Best post-holiday memory prompt."),
        (24, "Best balanced easyJet/easy-family introduction."),
        (25, "Best concise trust clarification."),
        (31, "Best travel-led Apple interaction."),
        (36, "Best travel-led Samsung introduction."),
        (41, "Best owner-value message before a flight."),
        (49, "Best trust-first re-entry for a resolved contact."),
    ]
    for number_value, reason in shortlist:
        p = doc.add_paragraph(style="List Number")
        set_spacing(p, after=5, line=1.15)
        add_text(p, f"Option {number_value:02d}: {reason}", size=10.2, color=INK)

    p = doc.add_paragraph()
    set_spacing(p, after=0, line=1.18)
    set_paragraph_box(p, BLUE_PALE, "C5DBEA")
    add_text(p, "NEXT DECISION", size=9.3, color=GREEN_DARK, bold=True)
    p.add_run().add_break()
    p.add_run().add_break()
    add_text(p, "Approve promising copy by option number first. Then map each approved option to an eligible cohort, timing rule, verified landing page, Meta template and controlled test size.", size=10.1, color=INK)

    core = doc.core_properties
    core.title = "easyTag WhatsApp Message Approval Catalogue"
    core.subject = "Numbered creative concepts and targeting conditions"
    core.author = "easyTag"
    core.keywords = "WhatsApp, messaging, CTR, easyTag, approval"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
